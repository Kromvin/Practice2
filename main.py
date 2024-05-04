import tkinter as tk
import psycopg2
from docx import Document
from docx.shared import Cm
DB_PARAMS = {
    "user": "veised1", 
    "password": "qweqwe", 
    "host": "localhost",
    "port": "5432",
    "database": "postgres"
}

def check_credentials():
    username = username_entry.get()
    password = password_entry.get()

    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("SELECT * FROM users WHERE login = %s AND password = %s;", (username, password))
        user = cursor.fetchone()

        if user:
            user_id, _, _, role = user
            if role == "admin":
                show_admin_panel()
            elif role == "student":
                show_user_panel(user_id, username)
            else:
                show_practicmng_panel()
        else:
            show_message("Неверные логин или пароль")

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при подключении к PostgreSQL: {error}")

    finally:
        if connection:
            connection.close()

def show_admin_panel():
    admin_window = tk.Toplevel(app)
    admin_window.title("Панель администратора")

    add_user_button = tk.Button(admin_window, text="Добавить пользователя", command=show_add_user_panel)
    add_user_button.grid(row=0, column=0, padx=5, pady=5)

    delete_user_button = tk.Button(admin_window, text="Удалить пользователя", command=show_delete_user_panel)
    delete_user_button.grid(row=1, column=0, padx=5, pady=5)

    admin_documents_button = tk.Button(admin_window, text="Документы", command=admin_documents)
    admin_documents_button.grid(row=2, column=0, padx=5, pady=5)

def show_add_user_panel():
    add_user_window = tk.Toplevel(app)
    add_user_window.title("Добавление пользователя")

    new_username_label = tk.Label(add_user_window, text="Логин:")
    new_username_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
    new_username_entry = tk.Entry(add_user_window)
    new_username_entry.grid(row=0, column=1, padx=5, pady=5)

    new_password_label = tk.Label(add_user_window, text="Пароль:")
    new_password_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")
    new_password_entry = tk.Entry(add_user_window, show="*")
    new_password_entry.grid(row=1, column=1, padx=5, pady=5)

    new_role_label = tk.Label(add_user_window, text="Роль:")
    new_role_label.grid(row=2, column=0, padx=5, pady=5, sticky="e")
    new_role_entry = tk.Entry(add_user_window)
    new_role_entry.grid(row=2, column=1, padx=5, pady=5)

    add_user_button = tk.Button(add_user_window, text="Добавить пользователя",
                            command=lambda: [add_user(new_username_entry.get(),
                                                      new_password_entry.get(), new_role_entry.get()),
                                            add_user_window.destroy()])
    add_user_button.grid(row=3, columnspan=2, padx=5, pady=5)

def show_practicmng_panel():
    practicmng_window = tk.Toplevel(app)
    practicmng_window.title("Панель практикующего")

    fill_info_button = tk.Button(practicmng_window, text="Заполнить информацию", command=lambda: add_practicmng_info())
    fill_info_button.grid(row=1, column=0, padx=5, pady=5)

def add_practicmng_info():
    add_practicmng_window = tk.Toplevel(app)
    add_practicmng_window.title("Добавить информацию о практикующем")

    dolgnost_label = tk.Label(add_practicmng_window, text="Должность:")
    dolgnost_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
    dolgnost_entry = tk.Entry(add_practicmng_window)
    dolgnost_entry.grid(row=0, column=1, padx=5, pady=5)

    secondname_label = tk.Label(add_practicmng_window, text="Фамилия:")
    secondname_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")
    secondname_entry = tk.Entry(add_practicmng_window)
    secondname_entry.grid(row=1, column=1, padx=5, pady=5)

    name_label = tk.Label(add_practicmng_window, text="Имя:")
    name_label.grid(row=2, column=0, padx=5, pady=5, sticky="e")
    name_entry = tk.Entry(add_practicmng_window)
    name_entry.grid(row=2, column=1, padx=5, pady=5)

    fathername_label = tk.Label(add_practicmng_window, text="Отчество:")
    fathername_label.grid(row=3, column=0, padx=5, pady=5, sticky="e")
    fathername_entry = tk.Entry(add_practicmng_window)
    fathername_entry.grid(row=3, column=1, padx=5, pady=5)

    adressProv_label = tk.Label(add_practicmng_window, text="Адрес проведения:")
    adressProv_label.grid(row=4, column=0, padx=5, pady=5, sticky="e")
    adressProv_entry = tk.Entry(add_practicmng_window)
    adressProv_entry.grid(row=4, column=1, padx=5, pady=5)

    gorodProv_label = tk.Label(add_practicmng_window, text="Город проведения:")
    gorodProv_label.grid(row=5, column=0, padx=5, pady=5, sticky="e")
    gorodProv_entry = tk.Entry(add_practicmng_window)
    gorodProv_entry.grid(row=5, column=1, padx=5, pady=5)

    save_button = tk.Button(add_practicmng_window, text="Сохранить", command=lambda: save_practicmng_info(
        dolgnost_entry.get(), secondname_entry.get(), name_entry.get(), fathername_entry.get(),
        adressProv_entry.get(), gorodProv_entry.get()))
    save_button.grid(row=6, columnspan=2, padx=5, pady=5)

def save_practicmng_info(dolgnost, secondname, name, fathername, adressProv, gorodProv):
    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("INSERT INTO practicmng (dolgnost, secondname, name, fathername, adressProv, gorodProv) "
                       "VALUES (%s, %s, %s, %s, %s, %s);",
                       (dolgnost, secondname, name, fathername, adressProv, gorodProv))
        connection.commit()

        show_message("Информация о практикующем успешно добавлена!")

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при добавлении информации о практикующем: {error}")

    finally:
        if connection:
            connection.close()

def add_user(username, password, role):
    connection = None

    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("SELECT MAX(id) FROM users;")
        max_id = cursor.fetchone()[0]  

        if max_id is None:
            new_id = 1
        else:
            new_id = max_id + 1

        cursor.execute("INSERT INTO users (id, login, password, role) VALUES (%s, %s, %s, %s);",
                       (new_id, username, password, role))
        connection.commit()

        show_message("Пользователь успешно добавлен!")

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при добавлении пользователя: {error}")

    finally:
        if connection:
            connection.close()

def show_delete_user_panel():
    delete_user_window = tk.Toplevel(app)
    delete_user_window.title("Удаление пользователя")

    username_label = tk.Label(delete_user_window, text="Логин пользователя:")
    username_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
    username_entry = tk.Entry(delete_user_window)
    username_entry.grid(row=0, column=1, padx=5, pady=5)

    delete_user_button = tk.Button(delete_user_window, text="Удалить пользователя",
                               command=lambda: [delete_user(username_entry.get()),
                                               delete_user_window.destroy()])
    delete_user_button.grid(row=1, columnspan=2, padx=5, pady=5)

def delete_user(username):
    connection = None

    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("SELECT * FROM users WHERE login = %s;", (username,))
        user = cursor.fetchone()

        if user:
            cursor.execute("DELETE FROM users WHERE login = %s;", (username,))
            connection.commit()
            show_message("Пользователь успешно удален!")
        else:
            show_message("Пользователь не найден.")

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при удалении пользователя: {error}")

    finally:
        if connection:
            connection.close()    

def admin_documents():
    admin_documents_window = tk.Toplevel(app)
    admin_documents_window.title("Документы администратора")

    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("SELECT id, name FROM doc;")
        documents = cursor.fetchall()

        for doc_id, doc_name in documents:
            doc_button = tk.Button(admin_documents_window, text=doc_name, command=lambda d=doc_id: open_document_admin(d))
            doc_button.pack(padx=5, pady=5)

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при получении документов: {error}")

    finally:
        if connection:
            connection.close()
    create_document_button = tk.Button(admin_documents_window, text="Создать новый документ", command=create_document_window)
    create_document_button.pack(padx=10, pady=10)

def open_document_admin(document_id):
    admin_documents_window = tk.Toplevel(app)
    admin_documents_window.title(f"Документ номер {document_id}")

    institute_label = tk.Label(admin_documents_window, text="Институт:")
    institute_label.grid(row=0, column=0, padx=5, pady=5)
    institute_entry = tk.Entry(admin_documents_window)
    institute_entry.grid(row=0, column=1, padx=5, pady=5)

    mestopract_label = tk.Label(admin_documents_window, text="Место проведения:")
    mestopract_label.grid(row=1, column=0, padx=5, pady=5)
    mestopract_entry = tk.Entry(admin_documents_window)
    mestopract_entry.grid(row=1, column=1, padx=5, pady=5)

    practice_type_label = tk.Label(admin_documents_window, text="Вид практики:")
    practice_type_label.grid(row=2, column=0, padx=5, pady=5)
    practice_type_entry = tk.Entry(admin_documents_window)
    practice_type_entry.grid(row=2, column=1, padx=5, pady=5)

    practice_kind_label = tk.Label(admin_documents_window, text="Тип практики:")
    practice_kind_label.grid(row=3, column=0, padx=5, pady=5)
    practice_kind_entry = tk.Entry(admin_documents_window)
    practice_kind_entry.grid(row=3, column=1, padx=5, pady=5)

    practice_god_label = tk.Label(admin_documents_window, text="Год проведения практики:")
    practice_god_label.grid(row=4, column=0, padx=5, pady=5)
    practice_god_entry = tk.Entry(admin_documents_window)
    practice_god_entry.grid(row=4, column=1, padx=5, pady=5)

    practicmng_label = tk.Label(admin_documents_window, text="Фамилия практикующего:")
    practicmng_label.grid(row=5, column=0, padx=5, pady=5)
    practicmng_entry = tk.StringVar(admin_documents_window)
    practicmng_list = ["Змеев"]  
    practicmng_entry.set(practicmng_list[0]) 
    practicmng_menu = tk.OptionMenu(admin_documents_window, practicmng_entry, *practicmng_list)
    practicmng_menu.grid(row=5, column=1, padx=5, pady=5)

    start_date_label = tk.Label(admin_documents_window, text="Дата начала практики (ДД.ММ):")
    start_date_label.grid(row=6, column=0, padx=5, pady=5)
    start_date_entry = tk.Entry(admin_documents_window)
    start_date_entry.grid(row=6, column=1, padx=5, pady=5)

    end_date_label = tk.Label(admin_documents_window, text="Дата окончания практики (ДД.ММ):")
    end_date_label.grid(row=7, column=0, padx=5, pady=5)
    end_date_entry = tk.Entry(admin_documents_window)
    end_date_entry.grid(row=7, column=1, padx=5, pady=5)

    save_button = tk.Button(admin_documents_window, text="Сохранить", command=lambda: save_document1(document_id,
                                                                                                   institute_entry.get(),
                                                                                                   mestopract_entry.get(),
                                                                                                   practice_type_entry.get(),
                                                                                                   practice_kind_entry.get(),
                                                                                                   practice_god_entry.get(),
                                                                                                   practicmng_entry.get(),
                                                                                                   start_date_entry.get(),
                                                                                                   end_date_entry.get()))
    save_button.grid(row=8, columnspan=2, padx=5, pady=5)

def save_document1(document_id, institute,mestopract, practice_type, practice_kind,god, practicmng_surname, start_date, end_date):
    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("SELECT id FROM practicmng WHERE secondname = %s;", (practicmng_surname,))
        practicmng_id = cursor.fetchone()[0]

        cursor.execute("UPDATE doc SET institut = %s,mestopract = %s, vidpract = %s, typePract = %s,god = %s, id_practicmng = %s, dateStart = %s, dateFinish = %s WHERE id = %s;",
                       (institute,mestopract, practice_type, practice_kind,god, practicmng_id, start_date, end_date, document_id))
        connection.commit()

        show_message(f"Документ {document_id} успешно сохранен!")

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при сохранении документа: {error}")

    finally:
        if connection:
            connection.close()

def save_document(name, supervisor_lastname):
    connection = None

    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("SELECT id FROM practicmng WHERE secondname = %s;", (supervisor_lastname,))
        supervisor_id = cursor.fetchone()[0]

        cursor.execute("INSERT INTO doc (name, id_practicmng) VALUES (%s, %s);", (name, supervisor_id))
        connection.commit()

        show_message("Документ успешно сохранен!")

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при сохранении документа: {error}")

    finally:
        if connection:
            connection.close()

def create_document_window():
    document_window = tk.Toplevel(app)
    document_window.title("Создание документа")

    name_label = tk.Label(document_window, text="Название документа:")
    name_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
    name_entry = tk.Entry(document_window)
    name_entry.grid(row=0, column=1, padx=5, pady=5)

    supervisor_label = tk.Label(document_window, text="Фамилия руководителя от ЮГУ:")
    supervisor_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")
    supervisor_entry = tk.Entry(document_window)
    supervisor_entry.grid(row=1, column=1, padx=5, pady=5)

    save_button = tk.Button(document_window, text="Сохранить", 
                            command=lambda: save_document(name_entry.get(), supervisor_entry.get()))
    save_button.grid(row=2, columnspan=2, padx=5, pady=5)

def show_user_panel(user_id, username):
    user_window = tk.Toplevel(app)
    user_window.title("Панель пользователя")

    documents_button = tk.Button(user_window, text="Документы", command=lambda: user_documents(user_id))
    documents_button.grid(row=0, column=0, padx=5, pady=5)

    fill_info_button = tk.Button(user_window, text="Заполнить информацию", command=lambda: add_user_info(username))
    fill_info_button.grid(row=1, column=0, padx=5, pady=5)

def add_user_info(username):
    add_user_info_window = tk.Toplevel(app)
    add_user_info_window.title("Добавить информацию о себе")

    group_label = tk.Label(add_user_info_window, text="Группа:")
    group_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
    group_entry = tk.Entry(add_user_info_window)
    group_entry.grid(row=0, column=1, padx=5, pady=5)

    curs_label = tk.Label(add_user_info_window, text="Курс:")
    curs_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")
    curs_entry = tk.Entry(add_user_info_window)
    curs_entry.grid(row=1, column=1, padx=5, pady=5)

    secondname_label = tk.Label(add_user_info_window, text="Фамилия:")
    secondname_label.grid(row=2, column=0, padx=5, pady=5, sticky="e")
    secondname_entry = tk.Entry(add_user_info_window)
    secondname_entry.grid(row=2, column=1, padx=5, pady=5)

    name_label = tk.Label(add_user_info_window, text="Имя:")
    name_label.grid(row=3, column=0, padx=5, pady=5, sticky="e")
    name_entry = tk.Entry(add_user_info_window)
    name_entry.grid(row=3, column=1, padx=5, pady=5)

    fathername_label = tk.Label(add_user_info_window, text="Отчество:")
    fathername_label.grid(row=4, column=0, padx=5, pady=5, sticky="e")
    fathername_entry = tk.Entry(add_user_info_window)
    fathername_entry.grid(row=4, column=1, padx=5, pady=5)

    napravlen_label = tk.Label(add_user_info_window, text="Направление подготовки:")
    napravlen_label.grid(row=5, column=0, padx=5, pady=5, sticky="e")
    napravlen_entry = tk.Entry(add_user_info_window)
    napravlen_entry.grid(row=5, column=1, padx=5, pady=5)

    save_button = tk.Button(add_user_info_window, text="Сохранить", command=lambda: save_user_info(username,
                                                                                                group_entry.get(),
                                                                                                 curs_entry.get(),
                                                                                                 secondname_entry.get(),
                                                                                                 name_entry.get(),
                                                                                                 fathername_entry.get(),
                                                                                                 napravlen_entry.get()))
    save_button.grid(row=6, columnspan=2, padx=5, pady=5)

def save_user_info(username,   group, curs, secondname, name, fathername, napravlen):
    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("SELECT id FROM users WHERE login = %s;", (username,))
        user_id = cursor.fetchone()[0]

        cursor.execute("INSERT INTO students (groupа, curs, secondname, name, fathername, napravlen, id_user) "
                       "VALUES (%s, %s, %s, %s, %s, %s, %s);",
                       (group, curs, secondname, name, fathername, napravlen, user_id))
        
        connection.commit()
        show_message("Информация успешно добавлена!")

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при добавлении информации: {error}")

    finally:
        if connection:
            connection.close()

def user_documents(user_id):
    admin_documents_window = tk.Toplevel(app)
    admin_documents_window.title("Документы студента")
    try:
        connection = psycopg2.connect(**DB_PARAMS)
        cursor = connection.cursor()

        cursor.execute("SELECT id, name FROM doc;")
        documents = cursor.fetchall()

        for doc_id, doc_name in documents:
            doc_button = tk.Button(admin_documents_window, text=doc_name + " student", command=lambda d=doc_id, u=user_id: open_document_user(d, u))
            doc_button.pack(padx=5, pady=5)

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при получении документов: {error}")

    finally:
        if connection:
            connection.close()

def open_document_user(document_id, user_id):
    try:
        with psycopg2.connect(**DB_PARAMS) as connection:
            with connection.cursor() as cursor:
                cursor.execute("SELECT institut, mestopract, vidpract, typePract, dateStart, dateFinish, god FROM doc WHERE id = %s;", (document_id,))
                document_info = cursor.fetchone()

        if document_info:
            document_window = tk.Toplevel(app)
            document_window.title(f"Документ номер {document_id}")

            labels = ["Институт:", "Место практики:", "Вид практики:", "Тип практики:", "Дата начала:", "Дата окончания:", "год проведения практики:", "Город проведения:"]
            for i, label_text in enumerate(labels):
                label = tk.Label(document_window, text=label_text)
                label.grid(row=i, column=0, padx=5, pady=5)
                if document_info and len(document_info) > i:
                    value_label = tk.Label(document_window, text=document_info[i])
                    value_label.grid(row=i, column=1, padx=5, pady=5)

            with psycopg2.connect(**DB_PARAMS) as connection_student:
                with connection_student.cursor() as cursor_student:
                    cursor_student.execute("SELECT groupа, curs, secondname, name, fathername, napravlen FROM students WHERE id_user = %s;", (user_id,))
                    student_info = cursor_student.fetchone()

            if student_info:
                student_labels = ["Группа:", "Курс:", "Фамилия студента:", "Имя студента:", "Отчество студента:", "Направление:"]
                for j, student_label_text in enumerate(student_labels):
                    label = tk.Label(document_window, text=student_label_text)
                    label.grid(row=len(labels) + j, column=0, padx=5, pady=5)
                    if len(student_info) > j:
                        value_label = tk.Label(document_window, text=student_info[j])
                        value_label.grid(row=len(labels) + j, column=1, padx=5, pady=5)

            with psycopg2.connect(**DB_PARAMS) as connection_practicmng:
                with connection_practicmng.cursor() as cursor_practicmng:
                    cursor_practicmng.execute("SELECT gorodProv FROM practicmng WHERE id = 1;")
                    gorodProv = cursor_practicmng.fetchone()[0]

                gorodProv_label = tk.Label(document_window, text=gorodProv)
                gorodProv_label.grid(row=len(labels) - 1, column=1, padx=5, pady=5)

            student_button = tk.Button(document_window, text="Создать документ", command=lambda: create_document_student((*student_info, gorodProv, *document_info)))
            student_button.grid(row=len(labels) + len(student_labels), columnspan=2, padx=5, pady=5)

        else:
            show_message("Документ не найден.")

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при открытии документа: {error}")

def create_document_student(document_data):
    try:
        input_window = tk.Toplevel(app)
        input_window.title("Добавление информации")

        fields = ["название документа", "оценка", "замечания", "проявленные качества", "выполнение индивидуального задания", "сложности", "выполненные задачи"]
        entry_labels = {}
        for i, field in enumerate(fields):
            label = tk.Label(input_window, text=f"{field.capitalize()}:")
            label.grid(row=i, column=0, padx=5, pady=5)
            entry = tk.Entry(input_window)
            entry.grid(row=i, column=1, padx=5, pady=5)
            entry_labels[field] = entry

        def save_data():
            try:
                additional_info = [entry_labels[field].get() for field in fields]
                with psycopg2.connect(**DB_PARAMS) as connection:
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT gorodProv FROM practicmng WHERE id = 1;")
                        gorodProv = cursor.fetchone()[0]
                        query = """
                            INSERT INTO studentdoc (groupа, curs, secondname, name, fathername, napravlen, gorodProv, 
                                                   institut, mestopract, vidpract, typePract, dateStart, dateFinish, 
                                                   god, namedoc, otsenka, zamech, kachestv, indzad, slojnosti, zadachi) 
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s);
                        """
                        data = document_data[:6] + document_data[6:] + tuple(additional_info)
                        print("Передаваемые данные:", data) 
                        cursor.execute(query, data)
                        connection.commit()
                        arr=get_studentdoc_data()
                        generate_report(arr)
                        
            except (Exception, psycopg2.Error) as error:
                show_message(f"Ошибка при сохранении данных: {error}")

        save_button = tk.Button(input_window, text="Сохранить", command=save_data)
        save_button.grid(row=len(fields), columnspan=2, padx=5, pady=5)

    except (Exception, psycopg2.Error) as error:
        show_message(f"Ошибка при создании документа: {error}")

def show_message(message):

    result_window = tk.Toplevel(app)
    result_window.title("Сообщение")
    result_label = tk.Label(result_window, text=message)
    result_label.pack(padx=10, pady=10)


def get_studentdoc_data():
    try:
        with psycopg2.connect(**DB_PARAMS) as connection:
            with connection.cursor() as cursor:
                cursor.execute("SELECT namedoc,mestopract,vidpract,typepract,institut,datestart,datefinish,god,groupа,curs,secondname,name,fathername,napravlen,otsenka,gorodprov,zamech,kachestv,indzad,slojnosti,zadachi FROM studentdoc WHERE id = 1")
                studentdoc_data = cursor.fetchall()
                cursor.execute("SELECT secondname,name,fathername FROM practicmng WHERE id = 1")
                practic_data=cursor.fetchall()
                for item in practic_data:
                    studentdoc_data.append(item)
        return studentdoc_data
    except (Exception, psycopg2.Error) as error:
        print(f"Ошибка при получении данных из таблицы studentdoc: {error}")
        return None


def generate_report(arr):
    doc = Document()
    doc.add_paragraph('МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ').alignment = 1
    doc.add_paragraph('РОССИЙСКОЙ ФЕДЕРАЦИИ').alignment = 1
    doc.add_paragraph(arr[0][4]).alignment = 1
    doc.add_paragraph('').alignment = 1
    doc.add_paragraph('').alignment = 1
    doc.add_paragraph('ОТЧЕТ').alignment = 1
    doc.add_paragraph('по производственной практике').alignment = 1
    doc.add_paragraph('«Учебная практика»').alignment = 1
    doc.add_paragraph('').alignment = 1
    doc.add_paragraph('').alignment = 1
    doc.add_paragraph('Место практики\t'+ arr[0][4])
    doc.add_paragraph('Студента\t'+ str(arr[0][9]) +'\tкурса    группы\t'+arr[0][8]+'\t'+ arr[0][10]+ arr[0][11]+ arr[0][12])
    doc.add_paragraph('Руководитель практики \t\t\t\t\t/' + arr[1][0]+ ' '+ arr[1][1][0]+ '.' + arr[1][2][0] + '.')
    doc.add_paragraph('от ЮГУ')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph(arr[0][15]+', ' + str(arr[0][7])).alignment = 1
    doc.add_page_break()



    doc.add_paragraph('ДНЕВНИК ПО  ' +arr[0][2][:-2].upper() +'ОЙ ПРАКТИКЕ').alignment = 1
    doc.add_paragraph('1.    Общие сведения').alignment = 1
    doc.add_paragraph('Фамилия, имя, отчество студента: '+ arr[0][10]+ arr[0][11]+ arr[0][12])
    doc.add_paragraph('Институт: '+ arr[0][4] )
    doc.add_paragraph('Курс '+ str(arr[0][9])+' Группа № '+arr[0][8])
    doc.add_paragraph('Место практики '+ arr[0][1])
    doc.add_paragraph('Сроки практики '+ arr[0][5]+'.'+str(arr[0][7])+'-'+ arr[0][6]+'.'+str(arr[0][7]))
    doc.add_paragraph('2. Производственная работа').alignment = 1
    table=doc.add_table(rows=7, cols=3)
    table.style='Table Grid'
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0,1).text='Выполненные виды работ в рамках задач (мероприятий),входящих в задание студента на практику'
    table.cell(0,2).text='Подпись руководителя практики от организации'
    table.cell(1,0).text='Дата'
    table.cell(1,1).text='Наименование работы'
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('Руководитель практики от организации\t'+arr[1][0]+ ' '+ arr[1][1][0]+ '.' + arr[1][2][0] + '.'+'/______________________________/')
    doc.add_page_break()



    doc.add_paragraph('Федеральное государственное бюджетное образовательное учреждение').alignment = 1
    doc.add_paragraph('высшего образования').alignment = 1
    doc.add_paragraph(arr[0][4]).alignment = 1
    doc.add_paragraph('Направление подготовки ' + arr[0][13]).alignment = 1
    doc.add_paragraph('ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ').alignment = 1
    doc.add_paragraph('на '+ arr[0][3][:-2] +'ую практику').alignment = 1
    doc.add_paragraph('Выдано\t'+ arr[0][10]+ arr[0][11]+ arr[0][12])
    doc.add_paragraph('Студенту\t'+str(arr[0][9])+'\tкурса   учебная группа\t№'+arr[0][8])
    doc.add_paragraph('Место прохождения практики:')
    doc.add_paragraph(arr[0][4])
    doc.add_paragraph('Сроки прохождения практики: с '+arr[0][5]+'.2024 года  по  '+arr[0][5]+'.2024 года')
    doc.add_paragraph('СОВМЕСТНЫЙ ПЛАН ПРОВЕДЕНИЯ ПРАКТИКИ').alignment = 1
    table1 = doc.add_table(rows=8, cols=3)
    table1.style = 'Table Grid'
    for cell in table.columns[1].cells:
        cell.width = Cm(8)
    table1.cell(0, 0).text = '№ п/п'
    table1.cell(0, 1).text = 'Наименование раздела (этапа) практики'
    table1.cell(0, 2).text = 'Срок прохождения этапа (периода) практики'
    table1.cell(1, 0).text = '1.'
    table1.cell(2, 0).text = '1.1.'
    table1.cell(3, 0).text = '2.'
    table1.cell(4, 0).text = '2.1.'
    table1.cell(5, 0).text = '3.'
    table1.cell(6, 0).text = '3.1.'
    table1.cell(5, 2).text = '2 ч.'
    table1.cell(5, 2).text = '216ч.'
    table1.cell(1, 1).text = 'Подготовительный этап'
    table1.cell(2, 1).text = 'Общее собрание, распределение по местам практики, инструктаж по технике безопасности'
    table1.cell(3, 1).text = 'Основной этап'
    table1.cell(4, 1).text = 'Выполнение производственных заданий, выполнение индивидуального задания, обработка и анализ полученной информации подготовка отчёта по практике, подготовка и оформление отчёта по практике.'
    table1.cell(5, 1).text = 'Заключительный этап'
    table1.cell(6, 1).text = 'Оформление отчета и дневника практики, подготовка презентации и защита отчета по практике '
    table1.cell(7, 1).text = 'ИТОГО'
    table1.cell(1, 2).text = '1ч.'
    table1.cell(3, 2).text = '213ч.'
    doc.add_paragraph('Руководители практики:')
    doc.add_paragraph('от ЮГУ\t\t\t\t_______________\t'+ arr[1][0]+ ' '+ arr[1][1][0]+ '.' + arr[1][2][0] + '.')
    doc.add_paragraph('от профильной организации   \t_______________\t'+ arr[1][0]+ ' '+ arr[1][1][0]+ '.' + arr[1][2][0] + '.')
    doc.add_paragraph('Задание принято к исполнению\t_______________\t'+ arr[0][10]+ arr[0][11]+ arr[0][12])


    doc.add_paragraph('Характеристика').alignment = 1
    doc.add_paragraph(arr[0][10]+ arr[0][11]+ arr[0][12]).alignment = 1
    doc.add_paragraph('студента курса '+str(arr[0][9])+' группы '+arr[0][8]).alignment = 1
    doc.add_paragraph('направления подготовки '+ arr[0][13]).alignment = 1
    doc.add_paragraph('ФГБОУ ВО «'+arr[0][4]+'»').alignment = 1
    doc.add_paragraph('Студент '+ arr[0][10]+ arr[0][11]+ arr[0][12]+' в период '+ arr[0][5]+'.'+str(arr[0][7])+'-'+ arr[0][6]+'.'+str(arr[0][7])+' проходил  '+arr[0][2][:-2]+'ую практику в '+arr[0][4]+', расположенной по адресу: '+arr[0][1]+'.').alignment = 1
    doc.add_paragraph('В ходе практики студент выполнял следующие задачи:')
    doc.add_paragraph(arr[0][20])
    doc.add_paragraph('В ходе выполнения практики продемонстрировал следующие качества '+arr[0][16]+'. С возникающими при работе проблемами справлялся '+arr[0][17]+'. Индивидуальное задание, предусмотренное программой практики, выполнено '+arr[0][18]+'.')
    doc.add_paragraph('Замечания '+arr[0][19])
    doc.add_paragraph('Работа '+ arr[0][10]+ arr[0][11]+ arr[0][12]+' оценивается на «'+arr[0][14]+'»')
    doc.add_paragraph('Руководитель практики')
    doc.add_paragraph('от предприятия\t\t\t\t______________/'+ arr[1][0]+ ' '+ arr[1][1][0]+ '.' + arr[1][2][0] + '.')
    doc.save(arr[0][0]+'.docx')

app = tk.Tk()
app.title("Проверка учетных данных")

username_label = tk.Label(app, text="Логин:")
username_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
username_entry = tk.Entry(app)
username_entry.grid(row=0, column=1, padx=5, pady=5)

password_label = tk.Label(app, text="Пароль:")
password_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")
password_entry = tk.Entry(app, show="*")
password_entry.grid(row=1, column=1, padx=5, pady=5)

login_button = tk.Button(app, text="Войти", command=check_credentials)
login_button.grid(row=2, columnspan=2, padx=5, pady=5)

app.mainloop()